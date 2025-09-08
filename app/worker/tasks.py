from app.worker.celery_app import celery
from app.db.session import SessionLocal
from app.core.config import settings
from app import crud
import minio
from PyPDF2 import PdfFileReader
from docx import Document
import io

minio_client = minio.Minio(
    settings.MINIO_ENDPOINT,
    access_key=settings.MINIO_ACCESS_KEY,
    secret_key=settings.MINIO_SECRET_KEY,
    secure=False,
)

@celery.task
def extract_metadata(file_id: int):
    db = SessionLocal()
    file = crud.file.get(db, id=file_id)
    if not file:
        return

    try:
        response = minio_client.get_object(settings.MINIO_BUCKET, file.s3_key)
        file_data = io.BytesIO(response.read())
    finally:
        response.close()
        response.release_conn()

    if file.filename.endswith(".pdf"):
        pdf = PdfFileReader(file_data)
        info = pdf.getDocumentInfo()
        file.pages = pdf.getNumPages()
        file.author = info.author
        file.title = info.title
        file.creator_tool = info.creator
    elif file.filename.endswith(".docx"):
        doc = Document(file_data)
        file.paragraphs = len(doc.paragraphs)
        file.tables = len(doc.tables)
        # docx files do not have standard metadata fields like author or title in the same way as PDFs
        # so we will leave them empty for now.

    db.add(file)
    db.commit()
    db.close()
